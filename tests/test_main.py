# ---------- TESTS FOR MAIN PIPELINE ORCHESTRATION ----------

import pytest
from unittest.mock import patch, MagicMock
from datetime import datetime
from docx import Document

from jobsai.main import main, pipeline_step
from jobsai.utils.exceptions import CancellationError


# Mock form submissions
mock_form_submissions = {
    "general": [
        {"job-level": ["Expert"]},
        {"job-boards": ["Duunitori"]},
        {"deep-mode": "Yes"},
        {"cover-letter-num": 1},
        {"cover-letter-style": ["Professional"]},
    ],
    "languages": [{"python": 5}, {"javascript": 3}],
    "additional-info": [{"additional-info": "Experienced developer"}],
}

# Mock agent responses
mock_profile = "Experienced Python developer with 5 years of experience."
mock_keywords = ["python developer", "software engineer"]
mock_jobs = [
    {
        "title": "Python Developer",
        "company": "Company A",
        "location": "Helsinki",
        "url": "https://example.com/job1",
        "description_snippet": "Python developer needed",
    }
]
mock_scored_jobs = [
    {
        "title": "Python Developer",
        "company": "Company A",
        "location": "Helsinki",
        "url": "https://example.com/job1",
        "score": 85,
        "matched_skills": ["Python"],
        "missing_skills": ["Docker"],
    }
]
mock_job_analysis = """
Job Analysis
========================================
Top 1 Jobs:

Title: Python Developer
Company: Company A
Location: Helsinki
Score: 85%
Instructions: Focus on Python experience.
----------------------------------------
"""
mock_document = Document()
mock_document.add_paragraph("Test cover letter")


@pytest.fixture
def mock_agents():
    """Create mock agent instances."""
    profiler = MagicMock()
    profiler.create_profile = MagicMock(return_value=mock_profile)

    query_builder = MagicMock()
    query_builder.create_keywords = MagicMock(return_value=mock_keywords)

    searcher = MagicMock()
    searcher.search_jobs = MagicMock(return_value=mock_jobs)

    scorer = MagicMock()
    scorer.score_jobs = MagicMock(return_value=mock_scored_jobs)

    analyzer = MagicMock()
    analyzer.write_analysis = MagicMock(return_value=mock_job_analysis)

    generator = MagicMock()
    generator.generate_letters = MagicMock(return_value=[mock_document])

    return {
        "profiler": profiler,
        "query_builder": query_builder,
        "searcher": searcher,
        "scorer": scorer,
        "analyzer": analyzer,
        "generator": generator,
    }


@patch("jobsai.main.GeneratorAgent")
@patch("jobsai.main.AnalyzerAgent")
@patch("jobsai.main.ScorerService")
@patch("jobsai.main.SearcherService")
@patch("jobsai.main.QueryBuilderAgent")
@patch("jobsai.main.ProfilerAgent")
@patch("jobsai.main.extract_form_data")
def test_main_pipeline_success(
    mock_extract_form_data,
    mock_profiler_class,
    mock_query_builder_class,
    mock_searcher_class,
    mock_scorer_class,
    mock_analyzer_class,
    mock_generator_class,
    mock_agents,
):
    """Test successful pipeline execution."""
    # Setup mocks
    mock_extract_form_data.return_value = {
        "job_boards": ["Duunitori"],
        "deep_mode": "Yes",
        "cover_letter_num": 1,
        "cover_letter_style": ["Professional"],
        "tech_stack": [[{"python": 5}]],
    }

    mock_profiler_class.return_value = mock_agents["profiler"]
    mock_query_builder_class.return_value = mock_agents["query_builder"]
    mock_searcher_class.return_value = mock_agents["searcher"]
    mock_scorer_class.return_value = mock_agents["scorer"]
    mock_analyzer_class.return_value = mock_agents["analyzer"]
    mock_generator_class.return_value = mock_agents["generator"]

    # Run pipeline
    result = main(mock_form_submissions)

    # Verify result structure
    assert "document" in result or "documents" in result
    assert "timestamp" in result
    assert "filename" in result or "filenames" in result

    # Verify all agents were called
    mock_agents["profiler"].create_profile.assert_called_once()
    mock_agents["query_builder"].create_keywords.assert_called_once()
    mock_agents["searcher"].search_jobs.assert_called_once()
    mock_agents["scorer"].score_jobs.assert_called_once()
    mock_agents["analyzer"].write_analysis.assert_called_once()
    mock_agents["generator"].generate_letters.assert_called_once()


@patch("jobsai.main.GeneratorAgent")
@patch("jobsai.main.AnalyzerAgent")
@patch("jobsai.main.ScorerService")
@patch("jobsai.main.SearcherService")
@patch("jobsai.main.QueryBuilderAgent")
@patch("jobsai.main.ProfilerAgent")
@patch("jobsai.main.extract_form_data")
def test_main_pipeline_progress_callback(
    mock_extract_form_data,
    mock_profiler_class,
    mock_query_builder_class,
    mock_searcher_class,
    mock_scorer_class,
    mock_analyzer_class,
    mock_generator_class,
    mock_agents,
):
    """Test that progress callback is called for each phase."""
    mock_extract_form_data.return_value = {
        "job_boards": ["Duunitori"],
        "deep_mode": "No",
        "cover_letter_num": 1,
        "cover_letter_style": ["Professional"],
        "tech_stack": [[{"python": 5}]],
    }

    mock_profiler_class.return_value = mock_agents["profiler"]
    mock_query_builder_class.return_value = mock_agents["query_builder"]
    mock_searcher_class.return_value = mock_agents["searcher"]
    mock_scorer_class.return_value = mock_agents["scorer"]
    mock_analyzer_class.return_value = mock_agents["analyzer"]
    mock_generator_class.return_value = mock_agents["generator"]

    progress_calls = []

    def progress_callback(phase, message):
        progress_calls.append((phase, message))

    main(mock_form_submissions, progress_callback=progress_callback)

    # Verify progress was called for each phase
    phases = [call[0] for call in progress_calls]
    assert "profiling" in phases
    assert "searching" in phases
    assert "scoring" in phases
    assert "analyzing" in phases
    assert "generating" in phases


@patch("jobsai.main.GeneratorAgent")
@patch("jobsai.main.AnalyzerAgent")
@patch("jobsai.main.ScorerService")
@patch("jobsai.main.SearcherService")
@patch("jobsai.main.QueryBuilderAgent")
@patch("jobsai.main.ProfilerAgent")
@patch("jobsai.main.extract_form_data")
def test_main_pipeline_cancellation(
    mock_extract_form_data,
    mock_profiler_class,
    mock_query_builder_class,
    mock_searcher_class,
    mock_scorer_class,
    mock_analyzer_class,
    mock_generator_class,
    mock_agents,
):
    """Test that cancellation check works."""
    mock_extract_form_data.return_value = {
        "job_boards": ["Duunitori"],
        "deep_mode": "No",
        "cover_letter_num": 1,
        "cover_letter_style": ["Professional"],
        "tech_stack": [[{"python": 5}]],
    }

    mock_profiler_class.return_value = mock_agents["profiler"]
    mock_query_builder_class.return_value = mock_agents["query_builder"]
    mock_searcher_class.return_value = mock_agents["searcher"]
    mock_scorer_class.return_value = mock_agents["scorer"]
    mock_analyzer_class.return_value = mock_agents["analyzer"]
    mock_generator_class.return_value = mock_agents["generator"]

    def cancellation_check():
        return True  # Always cancel

    with pytest.raises(CancellationError, match="cancelled before start"):
        main(mock_form_submissions, cancellation_check=cancellation_check)


@patch("jobsai.main.GeneratorAgent")
@patch("jobsai.main.AnalyzerAgent")
@patch("jobsai.main.ScorerService")
@patch("jobsai.main.SearcherService")
@patch("jobsai.main.QueryBuilderAgent")
@patch("jobsai.main.ProfilerAgent")
@patch("jobsai.main.extract_form_data")
def test_main_pipeline_multiple_cover_letters(
    mock_extract_form_data,
    mock_profiler_class,
    mock_query_builder_class,
    mock_searcher_class,
    mock_scorer_class,
    mock_analyzer_class,
    mock_generator_class,
    mock_agents,
):
    """Test pipeline with multiple cover letters."""
    mock_extract_form_data.return_value = {
        "job_boards": ["Duunitori"],
        "deep_mode": "No",
        "cover_letter_num": 2,
        "cover_letter_style": ["Professional"],
        "tech_stack": [[{"python": 5}]],
    }

    # Create multiple documents
    doc1 = Document()
    doc1.add_paragraph("Letter 1")
    doc2 = Document()
    doc2.add_paragraph("Letter 2")
    mock_agents["generator"].generate_letters.return_value = [doc1, doc2]

    mock_profiler_class.return_value = mock_agents["profiler"]
    mock_query_builder_class.return_value = mock_agents["query_builder"]
    mock_searcher_class.return_value = mock_agents["searcher"]
    mock_scorer_class.return_value = mock_agents["scorer"]
    mock_analyzer_class.return_value = mock_agents["analyzer"]
    mock_generator_class.return_value = mock_agents["generator"]

    result = main(mock_form_submissions)

    # Should return documents list format
    assert "documents" in result
    assert "filenames" in result
    assert len(result["documents"]) == 2
    assert len(result["filenames"]) == 2


@patch("jobsai.main.GeneratorAgent")
@patch("jobsai.main.AnalyzerAgent")
@patch("jobsai.main.ScorerService")
@patch("jobsai.main.SearcherService")
@patch("jobsai.main.QueryBuilderAgent")
@patch("jobsai.main.ProfilerAgent")
@patch("jobsai.main.extract_form_data")
def test_main_pipeline_error_handling(
    mock_extract_form_data,
    mock_profiler_class,
    mock_query_builder_class,
    mock_searcher_class,
    mock_scorer_class,
    mock_analyzer_class,
    mock_generator_class,
    mock_agents,
):
    """Test that pipeline errors are handled correctly."""
    mock_extract_form_data.return_value = {
        "job_boards": ["Duunitori"],
        "deep_mode": "No",
        "cover_letter_num": 1,
        "cover_letter_style": ["Professional"],
        "tech_stack": [[{"python": 5}]],
    }

    # Make profiler raise an error
    mock_agents["profiler"].create_profile.side_effect = Exception("LLM error")

    mock_profiler_class.return_value = mock_agents["profiler"]
    mock_query_builder_class.return_value = mock_agents["query_builder"]
    mock_searcher_class.return_value = mock_agents["searcher"]
    mock_scorer_class.return_value = mock_agents["scorer"]
    mock_analyzer_class.return_value = mock_agents["analyzer"]
    mock_generator_class.return_value = mock_agents["generator"]

    with pytest.raises(RuntimeError, match="Profiling candidate failed"):
        main(mock_form_submissions)


def test_pipeline_step_decorator_success():
    """Test that pipeline_step decorator logs correctly."""

    @pipeline_step("Test step", 1, 1)
    def test_function():
        return "success"

    result = test_function()
    assert result == "success"


def test_pipeline_step_decorator_error():
    """Test that pipeline_step decorator handles errors."""

    @pipeline_step("Test step", 1, 1)
    def test_function():
        raise ValueError("Test error")

    with pytest.raises(RuntimeError, match="Test step failed"):
        test_function()
