"""
Integration Tests for Complete Pipeline Flow.

Tests the complete pipeline execution from form submission to document generation,
including all agents and services working together.
"""

import pytest
from unittest.mock import patch, MagicMock, Mock
from datetime import datetime
from docx import Document

from jobsai.main import main
from jobsai.utils.exceptions import CancellationError


# Mock form submissions
mock_form_submissions = {
    "general": [
        {"job-level": ["Expert"]},
        {"job-boards": ["Duunitori", "Jobly"]},
        {"deep-mode": "Yes"},
        {"cover-letter-num": 2},
        {"cover-letter-style": ["Professional", "Friendly"]},
    ],
    "languages": [{"python": 5}, {"javascript": 3}],
    "databases": [{"postgresql": 4}],
    "additional-info": [{"additional-info": "Experienced developer"}],
}


@pytest.fixture
def mock_agents():
    """Create mock agent instances."""
    profiler = MagicMock()
    profiler.create_profile.return_value = (
        "Experienced Python developer with 5 years of experience."
    )

    query_builder = MagicMock()
    query_builder.create_keywords.return_value = [
        "python developer",
        "software engineer",
    ]

    searcher = MagicMock()
    searcher.search_jobs.return_value = [
        {
            "title": "Python Developer",
            "company": "Company A",
            "location": "Helsinki",
            "url": "https://example.com/job1",
            "description_snippet": "Python developer needed",
            "full_description": "We need a Python developer with experience.",
        }
    ]

    scorer = MagicMock()
    scorer.score_jobs.return_value = [
        {
            "title": "Python Developer",
            "company": "Company A",
            "location": "Helsinki",
            "url": "https://example.com/job1",
            "score": 85,
            "matched_skills": ["Python"],
            "missing_skills": ["Docker"],
        }
    ]

    analyzer = MagicMock()
    analyzer.write_analysis.return_value = """
Job Analysis
========================================
Top 1 Jobs:

Title: Python Developer
Company: Company A
Location: Helsinki
Score: 85%
Instructions: Focus on Python experience.
----------------------------------------
"""

    generator = MagicMock()
    doc1 = Document()
    doc1.add_paragraph("Cover letter 1")
    doc2 = Document()
    doc2.add_paragraph("Cover letter 2")
    generator.generate_letters.return_value = [doc1, doc2]

    return {
        "profiler": profiler,
        "query_builder": query_builder,
        "searcher": searcher,
        "scorer": scorer,
        "analyzer": analyzer,
        "generator": generator,
    }


@patch("jobsai.main.GeneratorAgent")
@patch("jobsai.main.AnalyzerAgent")
@patch("jobsai.main.ScorerService")
@patch("jobsai.main.SearcherService")
@patch("jobsai.main.QueryBuilderAgent")
@patch("jobsai.main.ProfilerAgent")
@patch("jobsai.main.extract_form_data")
def test_complete_pipeline_flow(
    mock_extract_form_data,
    mock_profiler_class,
    mock_query_builder_class,
    mock_searcher_class,
    mock_scorer_class,
    mock_analyzer_class,
    mock_generator_class,
    mock_agents,
):
    """Test complete pipeline flow from start to finish."""
    # Setup mocks
    mock_extract_form_data.return_value = {
        "job_boards": ["Duunitori", "Jobly"],
        "deep_mode": "Yes",
        "cover_letter_num": 2,
        "cover_letter_style": ["Professional", "Friendly"],
        "tech_stack": [[{"python": 5}, {"javascript": 3}, {"postgresql": 4}]],
    }

    mock_profiler_class.return_value = mock_agents["profiler"]
    mock_query_builder_class.return_value = mock_agents["query_builder"]
    mock_searcher_class.return_value = mock_agents["searcher"]
    mock_scorer_class.return_value = mock_agents["scorer"]
    mock_analyzer_class.return_value = mock_agents["analyzer"]
    mock_generator_class.return_value = mock_agents["generator"]

    # Run pipeline
    result = main(mock_form_submissions)

    # Verify result structure for multiple documents
    assert "documents" in result
    assert "timestamp" in result
    assert "filenames" in result
    assert len(result["documents"]) == 2
    assert len(result["filenames"]) == 2

    # Verify all agents were called in sequence
    mock_agents["profiler"].create_profile.assert_called_once()
    mock_agents["query_builder"].create_keywords.assert_called_once()
    mock_agents["searcher"].search_jobs.assert_called_once()
    mock_agents["scorer"].score_jobs.assert_called_once()
    mock_agents["analyzer"].write_analysis.assert_called_once()
    mock_agents["generator"].generate_letters.assert_called_once()


@patch("jobsai.main.GeneratorAgent")
@patch("jobsai.main.AnalyzerAgent")
@patch("jobsai.main.ScorerService")
@patch("jobsai.main.SearcherService")
@patch("jobsai.main.QueryBuilderAgent")
@patch("jobsai.main.ProfilerAgent")
@patch("jobsai.main.extract_form_data")
def test_pipeline_with_progress_callback(
    mock_extract_form_data,
    mock_profiler_class,
    mock_query_builder_class,
    mock_searcher_class,
    mock_scorer_class,
    mock_analyzer_class,
    mock_generator_class,
    mock_agents,
):
    """Test pipeline with progress callback."""
    mock_extract_form_data.return_value = {
        "job_boards": ["Duunitori"],
        "deep_mode": "Yes",
        "cover_letter_num": 1,
        "cover_letter_style": ["Professional"],
        "tech_stack": [[{"python": 5}]],
    }

    mock_profiler_class.return_value = mock_agents["profiler"]
    mock_query_builder_class.return_value = mock_agents["query_builder"]
    mock_searcher_class.return_value = mock_agents["searcher"]
    mock_scorer_class.return_value = mock_agents["scorer"]
    mock_analyzer_class.return_value = mock_agents["analyzer"]
    mock_generator_class.return_value = mock_agents["generator"]

    # Track progress updates
    progress_updates = []

    def progress_callback(phase, message):
        progress_updates.append((phase, message))

    # Run pipeline with progress callback
    result = main(mock_form_submissions, progress_callback=progress_callback)

    # Verify progress updates were called
    assert len(progress_updates) > 0
    assert any(phase == "profiling" for phase, _ in progress_updates)
    assert any(phase == "searching" for phase, _ in progress_updates)
    assert any(phase == "scoring" for phase, _ in progress_updates)
    assert any(phase == "analyzing" for phase, _ in progress_updates)
    assert any(phase == "generating" for phase, _ in progress_updates)

    # Verify result
    assert "document" in result or "documents" in result


@patch("jobsai.main.GeneratorAgent")
@patch("jobsai.main.AnalyzerAgent")
@patch("jobsai.main.ScorerService")
@patch("jobsai.main.SearcherService")
@patch("jobsai.main.QueryBuilderAgent")
@patch("jobsai.main.ProfilerAgent")
@patch("jobsai.main.extract_form_data")
def test_pipeline_cancellation(
    mock_extract_form_data,
    mock_profiler_class,
    mock_query_builder_class,
    mock_searcher_class,
    mock_scorer_class,
    mock_analyzer_class,
    mock_generator_class,
    mock_agents,
):
    """Test pipeline cancellation during execution."""
    mock_extract_form_data.return_value = {
        "job_boards": ["Duunitori"],
        "deep_mode": "Yes",
        "cover_letter_num": 1,
        "cover_letter_style": ["Professional"],
        "tech_stack": [[{"python": 5}]],
    }

    mock_profiler_class.return_value = mock_agents["profiler"]
    mock_query_builder_class.return_value = mock_agents["query_builder"]
    mock_searcher_class.return_value = mock_agents["searcher"]
    mock_scorer_class.return_value = mock_agents["scorer"]
    mock_analyzer_class.return_value = mock_agents["analyzer"]
    mock_generator_class.return_value = mock_agents["generator"]

    # Cancellation check that returns True after profiling
    call_count = [0]

    def cancellation_check():
        call_count[0] += 1
        # Cancel after first step
        return call_count[0] > 1

    # Run pipeline with cancellation
    with pytest.raises(CancellationError) as exc_info:
        main(mock_form_submissions, cancellation_check=cancellation_check)

    assert "cancelled" in str(exc_info.value).lower()

    # Profiler should have been called
    mock_agents["profiler"].create_profile.assert_called_once()
    # Generator should NOT have been called (cancelled before)
    mock_agents["generator"].generate_letters.assert_not_called()
