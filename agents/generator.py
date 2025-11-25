# ---------- GENERATOR AGENT ----------

# generate_letters
# _build_system_prompt
# _build_system_prompt

import logging
from typing import Optional

from docx import Document

from utils.llms import call_llm
from utils.normalization import normalize_text

from config.schemas import SkillProfile

logger = logging.getLogger(__name__)


class GeneratorAgent:
    """
    Generate cover letter.
    """

    def __init__(self):
        """
        Construct the GeneratorAgent class.
        """

    def asd(self):
        doc = Document()

        # Header block (name + links + contact info)
        doc.add_heading("Joni Mäkinen", level=0)

        p = doc.add_paragraph()
        p.add_run("jonimakinen.com\n")
        p.add_run("linkedin.com/in/joni-daniel-makinen\n")
        p.add_run("github.com/jonidaniel\n\n")
        p.add_run("joni-makinen@live.fi\n")
        p.add_run("+358405882001\n\n")

        doc.add_paragraph("November 24, 2025\n")

        doc.add_paragraph("Hiring Team")
        doc.add_paragraph("Vuono Group\n\n")

        # Body paragraphs
        body = [
            "Dear Hiring Team,\n",
            "AI technology is something I’m super excited about and that’s why "
            "I’m now applying for the Associate AI & Software Engineer role you have open in Helsinki.\n",
            "You’re looking at a highly motivated and competitive Business Information Technologies student "
            "determined to master the latest agentic AI technologies and become fully AI-native. "
            "I possess strong programming skills, broad knowledge of the software production life cycle, "
            "and an analytical mindset. With a passion for and hands-on project experience on AI and LLMs, "
            "I’m eager to contribute to your team’s work on …\n",
            "There are many languages and technologies mentioned in the job listing that the candidates are "
            "expected to have skills or familiarity on. Since 2020, I’ve built numerous study-related and independent "
            "software development projects with appealing user interfaces and secure backend systems with database "
            "connections. Python and JavaScript have come beside me since the beginning, and I love Node.js development. "
            "I’ve also deployed applications in AWS Cloud since last June and cloud implementations are definitely "
            "one of my expertise fields. I have a fair amount of experience with React also, but to be honest, "
            "I never really liked frontend development that much.\n",
            "But the main piece of work that summarizes my aspirations and would like for you to take a look at "
            "is JobsAI. It’s an agentic AI system that I’ve been working on lately that’s designed to automate most of "
            "one’s job searching efforts by finding and evaluating relevant job listings from the web and ultimately "
            "generating résumés and cover letters for them, based on the user’s skill profile.\n",
            "JobsAI consists of 7 separate phases (plan, assess, search, score, report, generate, notify) following "
            "a prompt chaining architecture, some of them leveraging LLMs, some being deterministic. Based on the user’s "
            "initial skillsets and experiences input, the system creates a skill profile of the user, which is used to "
            "scrape relevant job listings off popular job boards (Duunitori etc.) in the web. The listings are then scored "
            "and organized, and an analysis is made of the whole process thus far. Then, tailored résumé and cover letter "
            "suggestions are drafted for the candidate to use in applying for the jobs. Finally, the analysis and the "
            "documents are sent to the user’s email, if required in the settings. The whole workflow is carried out from "
            "start to finish with only a single user text input, which shows that automation is really in the heart of JobsAI.\n",
            "The system leverages today’s most sophisticated agentic AI frameworks like OpenAI Agents SDK and LangChain. "
            "All LLM interactions are fine-tuned for highest-quality outcomes. JobsAI acts as a clear demonstration of my "
            "ability to design and implement complete end-to-end agentic systems that automate otherwise laborious processes "
            "and finally deliver end results that are of real value.\n",
            "Nowadays almost anyone can build an app in mere minutes. It’s clear that in the future software engineers are "
            "more like team leads instead of lone coders, controlling hordes of agents doing the actual coding work. "
            "That’s the future and the one thing I really want to learn about. I’m interested in developing agentic automation "
            "solutions that have a real impact on the world. The agentic landscape is so interesting to me I could spend 12 hours "
            "every day developing JobsAI.\n",
            "In the beginning I said I was an extremely motivated individual. I do learn quickly, and you saying you offer real "
            "responsibility from day one sounds great, and that’s going to send us in the right direction right from the start. "
            "I’ll dedicate myself fully to this role and for my future at Vuono Group.\n",
            "I would gladly welcome the opportunity to discuss how my skills, experience, and passion for AI and LLMs can aid "
            "Vuono Group in solving critical business challenges. Thank you.\n",
            "Best regards,\n",
            "Joni Mäkinen",
        ]

        for paragraph in body:
            doc.add_paragraph(paragraph)

        # Save the file
        doc.save("data/cover-letters/cover_letter.docx")

    # ------------------------------
    # Public interface
    # ------------------------------
    def generate_letters(
        self,
        skill_profile: SkillProfile,
        job_report: str,
        employer: Optional[str] = None,
        job_title: Optional[str] = None,
        style: str = "professional",
    ) -> str:
        """
        Produce a tailored job-application message based on
        the candidate's skills and the job report.

        Args:
            skill_profile:
            job_report:
            employer:
            job_title:
            style:

        Returns:
            output: the generated text
        """

        self.asd()

        system_prompt = self._build_system_prompt(style)
        user_prompt = self._build_user_prompt(
            skill_profile, job_report, employer, job_title
        )

        logger.info(" GENERATING APPLICATION TEXT...")

        raw = call_llm(system_prompt, user_prompt)

        output = normalize_text(raw)

        logger.info(" APPLICATION TEXT GENERATION COMPLETED\n")

        return output

    # ------------------------------
    # Internal functions
    # ------------------------------

    def _build_system_prompt(self, style: str) -> str:
        """
        Build system prompt.

        Args:
            style:
        """

        tone_instructions = {
            "professional": (
                "Write in a clear, respectful, concise, professional tone. "
                "Use well-structured paragraphs. Avoid exaggerations."
            ),
            "friendly": ("Write in a warm, positive tone but keep it professional."),
            "confident": (
                "Write with a confident, proactive tone without sounding arrogant."
            ),
        }

        base_style = tone_instructions.get(style, tone_instructions["professional"])

        return (
            "You are a professional cover letter writer. "
            "Your goal is to produce polished text suitable for real-world job applications.\n"
            "Follow this style:\n"
            f"{base_style}\n"
        )

    def _build_user_prompt(
        self,
        skill_profile: SkillProfile,
        job_report: str,
        employer: Optional[str],
        job_title: Optional[str],
    ) -> str:
        """
        Build user prompt.

        Args:
            skill_profile:
            job_report:
            employer:
            job_title:

        Returns:
            "
            Generate a tailored job-application message.

            Candidate Skill Profile (JSON):
            {skill_profile.model_dump_json(indent=2)}

            Job Match Analysis:
            {job_report}

            {employer_text}{title_text}

            Instructions:
            - Produce a compelling but concise job-application message.
            - Highlight the candidate's relevant skills based on the report.
            - If employer or job title are given, tailor the message to them.
            - Keep it truthful, specific, and readable.
            "
        """

        employer_text = f"Employer: {employer}\n" if employer else ""
        title_text = f"Target job title: {job_title}\n" if job_title else ""

        return f"""
            Generate a tailored job-application message.

            Candidate Skill Profile (JSON):
            {skill_profile.model_dump_json(indent=2)}

            Job Match Analysis:
            {job_report}

            {employer_text}{title_text}

            Instructions:
            - Produce a compelling but concise job-application message.
            - Highlight the candidate’s relevant skills based on the report.
            - If employer or job title are given, tailor the message to them.
            - Keep it truthful, specific, and readable.
            """
